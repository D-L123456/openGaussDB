import os
import logging

import docx
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.models.knowledge import KnowledgeNode
from app.core.docx_utils import natural_sort_key, extract_images_from_para, clean_figure_references, IMAGES_DIR

logger = logging.getLogger(__name__)


class KnowledgeTreeService:
    async def build_tree(self, docx_dir: str, db: AsyncSession) -> dict:
        os.makedirs(IMAGES_DIR, exist_ok=True)
        files = sorted([f for f in os.listdir(docx_dir) if f.endswith(".docx")], key=natural_sort_key)
        total_nodes = 0

        for filename in files:
            filepath = os.path.join(docx_dir, filename)
            doc = docx.Document(filepath)
            chapter_name = filename.replace(".docx", "")

            chapter_content_parts: list[str] = []
            current_h2_id = None
            current_h2_title = ""
            current_h2_parts: list[str] = []
            current_h3_id = None
            current_h3_title = ""
            current_h3_parts: list[str] = []

            async def flush_h3():
                nonlocal current_h3_id, total_nodes
                if current_h3_id and current_h3_parts:
                    content = "\n".join(current_h3_parts).strip()
                    if content:
                        existing = await db.get(KnowledgeNode, current_h3_id)
                        if existing:
                            existing.content = content

            chapter_node_id = await self._upsert_node(
                db=db,
                parent_id=None,
                chapter=chapter_name,
                section=chapter_name,
                title=chapter_name,
                content="",
                sort_order=total_nodes,
            )
            total_nodes += 1

            for para_idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                image_refs = extract_images_from_para(para, doc, chapter_name, para_idx, use_markdown=True)
                style_name = para.style.name if para.style else ""

                if "Heading 2" in style_name:
                    await flush_h3()
                    current_h3_id = None
                    current_h3_title = ""
                    current_h3_parts = []

                    current_h2_id = await self._upsert_node(
                        db=db,
                        parent_id=chapter_node_id,
                        chapter=chapter_name,
                        section=text,
                        title=text,
                        content="",
                        sort_order=total_nodes,
                    )
                    current_h2_title = text
                    current_h2_parts = []
                    total_nodes += 1

                elif "Heading 3" in style_name:
                    await flush_h3()

                    current_h3_id = await self._upsert_node(
                        db=db,
                        parent_id=current_h2_id or chapter_node_id,
                        chapter=chapter_name,
                        section=f"{current_h2_title} > {text}" if current_h2_title else text,
                        title=text,
                        content="",
                        sort_order=total_nodes,
                    )
                    current_h3_title = text
                    current_h3_parts = []
                    total_nodes += 1

                else:
                    parts_to_add = []
                    if image_refs:
                        parts_to_add.extend(image_refs)
                    if text:
                        cleaned = clean_figure_references(text)
                        if cleaned:
                            parts_to_add.append(cleaned)
                    if not parts_to_add:
                        continue

                    if current_h3_id:
                        current_h3_parts.extend(parts_to_add)
                    elif current_h2_id:
                        current_h2_parts.extend(parts_to_add)
                    else:
                        chapter_content_parts.extend(parts_to_add)

            await flush_h3()

            if current_h2_id and current_h2_parts:
                h2_node = await db.get(KnowledgeNode, current_h2_id)
                if h2_node:
                    existing_content = h2_node.content or ""
                    new_content = "\n".join(current_h2_parts).strip()
                    h2_node.content = existing_content + "\n" + new_content if existing_content else new_content

            if chapter_content_parts:
                ch_node = await db.get(KnowledgeNode, chapter_node_id)
                if ch_node:
                    ch_node.content = "\n".join(chapter_content_parts).strip()

            for table in doc.tables:
                table_text = self._extract_table(table)
                if table_text.strip() and current_h3_id:
                    current_h3_parts.append(table_text)
                elif table_text.strip() and current_h2_id:
                    current_h2_parts.append(table_text)

        await db.commit()
        return {"total_nodes": total_nodes, "status": "success"}

    def _extract_table(self, table) -> str:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)

    async def _upsert_node(
        self,
        db: AsyncSession,
        parent_id: str | None,
        chapter: str,
        section: str,
        title: str,
        content: str,
        sort_order: int,
    ) -> str:
        node = KnowledgeNode(
            parent_id=parent_id,
            chapter=chapter,
            section=section,
            title=title,
            content=content,
            sort_order=sort_order,
        )
        db.add(node)
        await db.flush()
        return str(node.id)

    async def get_tree(self, db: AsyncSession) -> list[dict]:
        result = await db.execute(
            select(KnowledgeNode).order_by(KnowledgeNode.sort_order)
        )
        all_nodes = result.scalars().all()

        node_map = {}
        root_nodes = []

        for node in all_nodes:
            node_data = {
                "id": str(node.id),
                "parent_id": str(node.parent_id) if node.parent_id else None,
                "chapter": node.chapter,
                "section": node.section,
                "title": node.title,
                "content": node.content,
                "sort_order": node.sort_order,
                "children": [],
            }
            node_map[str(node.id)] = node_data

        for node_data in node_map.values():
            parent_id = node_data["parent_id"]
            if parent_id and parent_id in node_map:
                node_map[parent_id]["children"].append(node_data)
            else:
                root_nodes.append(node_data)

        return root_nodes

    async def search_knowledge(self, query: str, top_k: int = 5) -> list[dict]:
        from app.services.vector_store import vector_store
        results = await vector_store.search(query, top_k=top_k)
        search_results = []
        for r in results:
            meta = r["metadata"]
            search_results.append({
                "node_id": "",
                "chapter": meta.get("chapter", ""),
                "section": meta.get("section", ""),
                "title": meta.get("title", ""),
                "content": r["content"],
                "score": r["score"],
            })
        return search_results


knowledge_tree_service = KnowledgeTreeService()
