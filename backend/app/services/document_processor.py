import os

import docx
from app.core.config import settings
from app.core.docx_utils import natural_sort_key, extract_images_from_para, clean_figure_references


class DocumentChunk:
    chapter: str
    section: str
    title: str
    content: str
    metadata: dict

    def __init__(self, chapter: str, section: str, title: str, content: str, metadata: dict):
        self.chapter = chapter
        self.section = section
        self.title = title
        self.content = content
        self.metadata = metadata


class DocumentProcessor:
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 100):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def extract_documents(self, docx_dir: str | None = None) -> list[DocumentChunk]:
        doc_dir = docx_dir or settings.docx_dir
        if not doc_dir:
            raise ValueError("未配置文档目录，请设置DOCX_DIR环境变量")

        chunks: list[DocumentChunk] = []
        files = sorted([f for f in os.listdir(doc_dir) if f.endswith(".docx")], key=natural_sort_key)

        for filename in files:
            filepath = os.path.join(doc_dir, filename)
            doc = docx.Document(filepath)
            chapter_name = filename.replace(".docx", "")
            file_chunks = self._parse_document(doc, chapter_name)
            chunks.extend(file_chunks)

        return chunks

    def _parse_document(self, doc: docx.Document, chapter_name: str) -> list[DocumentChunk]:
        chunks: list[DocumentChunk] = []
        current_section = ""
        current_title = ""
        current_content_parts: list[str] = []

        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            image_refs = extract_images_from_para(para, doc, chapter_name, para_idx, use_markdown=False)

            style_name = para.style.name if para.style else ""

            if "Heading 1" in style_name:
                if current_content_parts:
                    content = "\n".join(current_content_parts)
                    chunks.extend(self._split_content(
                        chapter_name, current_section, current_title, content
                    ))
                current_section = text
                current_title = text
                current_content_parts = []
            elif "Heading 2" in style_name:
                if current_content_parts:
                    content = "\n".join(current_content_parts)
                    chunks.extend(self._split_content(
                        chapter_name, current_section, current_title, content
                    ))
                current_section = text
                current_title = text
                current_content_parts = []
            elif "Heading 3" in style_name:
                if current_content_parts:
                    content = "\n".join(current_content_parts)
                    chunks.extend(self._split_content(
                        chapter_name, current_section, current_title, content
                    ))
                current_title = text
                current_content_parts = [text]
            else:
                parts_to_add = []
                if image_refs:
                    parts_to_add.extend(image_refs)
                if text:
                    cleaned = clean_figure_references(text)
                    if cleaned:
                        parts_to_add.append(cleaned)
                if parts_to_add:
                    current_content_parts.extend(parts_to_add)

        if current_content_parts:
            content = "\n".join(current_content_parts)
            chunks.extend(self._split_content(
                chapter_name, current_section, current_title, content
            ))

        for table in doc.tables:
            table_text = self._extract_table(table)
            if table_text.strip():
                chunks.append(DocumentChunk(
                    chapter=chapter_name,
                    section=current_section,
                    title=f"{current_title} - 表格",
                    content=table_text,
                    metadata={"type": "table"}
                ))

        return chunks

    def _extract_table(self, table) -> str:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)

    def _split_content(
        self, chapter: str, section: str, title: str, content: str
    ) -> list[DocumentChunk]:
        if len(content) <= self.chunk_size:
            return [DocumentChunk(
                chapter=chapter,
                section=section,
                title=title,
                content=content,
                metadata={"type": "text"}
            )]

        chunks = []
        start = 0
        while start < len(content):
            end = start + self.chunk_size
            chunk_text = content[start:end]

            if end < len(content):
                last_period = max(
                    chunk_text.rfind("。"),
                    chunk_text.rfind("；"),
                    chunk_text.rfind("\n"),
                )
                if last_period > self.chunk_size // 2:
                    chunk_text = chunk_text[: last_period + 1]
                    end = start + last_period + 1

            chunks.append(DocumentChunk(
                chapter=chapter,
                section=section,
                title=title,
                content=chunk_text.strip(),
                metadata={"type": "text", "chunk_index": len(chunks)}
            ))
            start = end - self.chunk_overlap

        return chunks
